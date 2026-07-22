import os
import uuid
import ffmpeg
import imageio_ffmpeg
from PIL import Image
from utils import get_temp_dir, cleanup_file

import re

def _run_ffmpeg_with_progress(stream, output_path, progress_callback):
    try:
        process = stream.run_async(
            cmd=imageio_ffmpeg.get_ffmpeg_exe(),
            pipe_stderr=True,
            pipe_stdout=False,
            quiet=True
        )
        
        size_pattern = re.compile(r"size=\s*(\d+[a-zA-Z]+)")
        speed_pattern = re.compile(r"speed=\s*([\d.]+x)")
        
        while True:
            line = process.stderr.readline()
            if not line:
                break
            
            line_str = line.decode('utf-8', errors='ignore')
            
            if progress_callback:
                size_match = size_pattern.search(line_str)
                speed_match = speed_pattern.search(line_str)
                
                if size_match and speed_match:
                    size = size_match.group(1)
                    speed = speed_match.group(1)
                    progress_callback(f"Converting... Size: {size}, Speed: {speed}")
                    
        process.wait()
        if process.returncode != 0:
            print(f"FFmpeg error with code {process.returncode}")
            return None
            
        return output_path
    except Exception as e:
        print(f"Error during ffmpeg execution: {e}")
        return None

def convert_video_to_audio(input_path, output_format='mp3', progress_callback=None):
    """
    Converts a video or audio file to a specified audio format.
    """
    temp_dir = get_temp_dir()
    output_filename = f"{uuid.uuid4()}.{output_format}"
    output_path = os.path.join(temp_dir, output_filename)
    
    stream = (
        ffmpeg
        .input(input_path)
        .output(output_path, acodec='libmp3lame' if output_format == 'mp3' else 'copy', qscale=2)
        .overwrite_output()
    )
    
    return _run_ffmpeg_with_progress(stream, output_path, progress_callback)

def convert_video_format(input_path, output_format='mp4', progress_callback=None):
    """
    Converts a video file to a different video format.
    """
    temp_dir = get_temp_dir()
    output_filename = f"{uuid.uuid4()}.{output_format}"
    output_path = os.path.join(temp_dir, output_filename)
    
    stream = (
        ffmpeg
        .input(input_path)
        .output(output_path)
        .overwrite_output()
    )
    
    return _run_ffmpeg_with_progress(stream, output_path, progress_callback)

def convert_image_format(input_path, output_format='png'):
    """
    Converts an image file to a different format using Pillow.
    """
    temp_dir = get_temp_dir()
    output_filename = f"{uuid.uuid4()}.{output_format}"
    output_path = os.path.join(temp_dir, output_filename)
    
    try:
        with Image.open(input_path) as img:
            rgb_im = img.convert('RGB')
            rgb_im.save(output_path, format=output_format.upper())
        return output_path
    except Exception as e:
        print(f"Error converting image format: {e}")
        return None

def convert_document_format(input_path, output_format='pdf'):
    """
    Converts a document (PDF, DOCX, TXT) to another document format (PDF, DOCX, TXT).
    """
    temp_dir = get_temp_dir()
    output_filename = f"{uuid.uuid4()}.{output_format}"
    output_path = os.path.join(temp_dir, output_filename)
    
    ext = os.path.splitext(input_path)[1].lower()
    extracted_text = ""
    
    try:
        if ext == '.pdf':
            import fitz
            doc = fitz.open(input_path)
            for page in doc:
                extracted_text += page.get_text() + "\n"
        elif ext in ['.docx', '.doc']:
            import docx
            doc = docx.Document(input_path)
            extracted_text = "\n".join([p.text for p in doc.paragraphs])
        else:
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                extracted_text = f.read()
                
        if not extracted_text.strip():
            extracted_text = "No text content found in original file."

        if output_format == 'txt':
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(extracted_text)
            return output_path
            
        elif output_format == 'docx':
            import docx
            doc = docx.Document()
            for line in extracted_text.split('\n'):
                doc.add_paragraph(line)
            doc.save(output_path)
            return output_path
            
        elif output_format == 'pdf':
            import fitz
            doc = fitz.open()
            page = doc.new_page()
            margin = 50
            rect = fitz.Rect(margin, margin, page.rect.width - margin, page.rect.height - margin)
            page.insert_textbox(rect, extracted_text, fontsize=11)
            doc.save(output_path)
            return output_path
            
    except Exception as e:
        print(f"Error converting document: {e}")
        return None

def translate_and_dub_media(input_path, target_lang='km', is_video=True, progress_callback=None):
    """
    Segment-level timed dubbing:
    1. Transcribes speech WITH timestamps (each phrase: start_sec, end_sec, text).
    2. Translates each segment individually.
    3. Generates TTS for each translated segment.
    4. Places each TTS clip at its exact original timestamp using FFmpeg adelay.
    5. Mixes timed dubbed audio with original background audio at 20% volume.
    6. Merges final audio track into the video.
    """
    import asyncio
    import subprocess
    temp_dir = get_temp_dir()

    from plugins.document_parser import transcribe_with_timestamps, transcribe_audio_video
    from plugins.ai_handler import get_ai_response
    from gtts import gTTS

    lang_names = {
        'km': 'Khmer', 'en': 'English', 'zh': 'Chinese (Mandarin)',
        'fr': 'French', 'es': 'Spanish', 'ja': 'Japanese',
        'ko': 'Korean', 'de': 'German', 'vi': 'Vietnamese',
        'th': 'Thai', 'id': 'Indonesian', 'ru': 'Russian', 'ar': 'Arabic'
    }
    target_lang_name = lang_names.get(target_lang[:2], 'Khmer')
    tts_lang = 'km' if target_lang.startswith('km') else ('zh-CN' if target_lang.startswith('zh') else target_lang[:2])

    ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
    orig_dur = get_video_duration(input_path)

    # ── Step 1: Timestamped transcription ────────────────────────────────────
    if progress_callback: progress_callback("⏳ Transcribing speech with timestamps...")
    timed_segments = transcribe_with_timestamps(input_path)

    # ── Step 2: Timed segment-level dubbing path ──────────────────────────────
    if timed_segments and len(timed_segments) >= 1:
        if progress_callback: progress_callback(f"🌐 Translating {len(timed_segments)} speech segments to {target_lang_name}...")

        # Batch-translate all segments at once (numbered list for efficiency)
        numbered_src = "\n".join(f"{i+1}. {seg[2]}" for i, seg in enumerate(timed_segments))
        prompt = (
            f"You are a professional movie dubbing translator.\n"
            f"Translate each numbered line into natural spoken {target_lang_name}.\n"
            f"Rules:\n"
            f"1. Keep the SAME numbering (1. 2. 3. ...).\n"
            f"2. Each translated line should be concise — similar length to the original spoken line.\n"
            f"3. Output ONLY the numbered translated lines — NO headers, NO explanations, NO markdown.\n\n"
            f"{numbered_src}"
        )
        try:
            try:
                loop = asyncio.get_running_loop()
                batch_translation = asyncio.run_coroutine_threadsafe(get_ai_response(0, prompt), loop).result(60)
            except RuntimeError:
                batch_translation = asyncio.run(get_ai_response(0, prompt))
        except Exception as e:
            print(f"Batch translation error: {e}")
            batch_translation = numbered_src  # fallback: use original text

        # Parse numbered lines from AI response
        translated_lines = {}
        for line in batch_translation.strip().splitlines():
            m = re.match(r'^(\d+)\.\s*(.+)', line.strip())
            if m:
                translated_lines[int(m.group(1))] = m.group(2).strip()

        if progress_callback: progress_callback("🎙 Generating timed voiceover segments (mouth-sync)...")

        # Generate TTS for each segment, speed-fit to slot, then adelay to start time
        tts_clips = []  # list of (delay_ms, fitted_tts_path)
        for i, (start_sec, end_sec, orig_text) in enumerate(timed_segments):
            seg_num = i + 1
            translated_text = translated_lines.get(seg_num, orig_text)
            translated_text = re.sub(r'[*#_~`>\[\]\(\)]', ' ', translated_text)
            translated_text = " ".join(translated_text.split()).strip()
            if not translated_text:
                continue

            slot_dur = max(0.3, end_sec - start_sec)  # original speech slot duration
            raw_seg = os.path.join(temp_dir, f"seg_raw_{i}_{uuid.uuid4().hex}.mp3")
            try:
                tts = gTTS(text=translated_text, lang=tts_lang, slow=False)
                tts.save(raw_seg)
                if not (os.path.exists(raw_seg) and os.path.getsize(raw_seg) > 0):
                    continue

                # ── Mouth-sync: speed-fit TTS to slot duration ──────────────
                tts_seg_dur = get_video_duration(raw_seg)
                fitted_seg = raw_seg
                if tts_seg_dur > 0 and tts_seg_dur != slot_dur:
                    speed_ratio = tts_seg_dur / slot_dur  # >1 means TTS is slower than slot
                    # Clamp: don't go faster than 2.0x or slower than 0.6x
                    speed_ratio = max(0.6, min(2.0, speed_ratio))
                    if abs(speed_ratio - 1.0) > 0.05:  # only adjust if >5% difference
                        fitted_path = os.path.join(temp_dir, f"seg_fit_{i}_{uuid.uuid4().hex}.mp3")
                        try:
                            # atempo supports 0.5–2.0; chain two filters for ratios > 2.0
                            if speed_ratio > 2.0:
                                af = f"atempo=2.0,atempo={speed_ratio/2.0:.4f}"
                            else:
                                af = f"atempo={speed_ratio:.4f}"
                            subprocess.run(
                                [ffmpeg_exe, "-y", "-i", raw_seg,
                                 "-af", af,
                                 "-t", str(slot_dur),  # hard-trim to slot length
                                 fitted_path],
                                capture_output=True, check=True
                            )
                            if os.path.exists(fitted_path) and os.path.getsize(fitted_path) > 0:
                                fitted_seg = fitted_path
                        except Exception as tempo_e:
                            print(f"atempo fit error seg {i}: {tempo_e}")

                delay_ms = max(0, int(start_sec * 1000))
                tts_clips.append((delay_ms, fitted_seg))
            except Exception as tts_e:
                print(f"TTS error for segment {i}: {tts_e}")

        if tts_clips:
            if progress_callback: progress_callback("🔀 Assembling mouth-synced dubbed audio track...")
            dubbed_audio_path = os.path.join(temp_dir, f"timed_dub_{uuid.uuid4().hex}.mp3")
            inputs_cmd = []
            filter_parts = []
            for idx, (delay_ms, tts_path) in enumerate(tts_clips):
                inputs_cmd += ["-i", tts_path]
                filter_parts.append(f"[{idx}:a]adelay={delay_ms}|{delay_ms}[a{idx}]")
            n = len(tts_clips)
            refs = "".join(f"[a{j}]" for j in range(n))
            filter_complex = ";".join(filter_parts) + f";{refs}amix=inputs={n}:duration=longest:dropout_transition=0[dub]"

            timed_ok = False
            try:
                cmd = [ffmpeg_exe, "-y"] + inputs_cmd + [
                    "-filter_complex", filter_complex,
                    "-map", "[dub]",
                    dubbed_audio_path
                ]
                if orig_dur > 0:
                    cmd = cmd[:-1] + ["-t", str(orig_dur), dubbed_audio_path]
                subprocess.run(cmd, capture_output=True, check=True)
                timed_ok = os.path.exists(dubbed_audio_path) and os.path.getsize(dubbed_audio_path) > 0
            except Exception as asm_e:
                print(f"Timed assembly error: {asm_e}")

            if timed_ok:
                # Pad assembled dub to full video length
                tts_dur = get_video_duration(dubbed_audio_path)
                if orig_dur > 0 and tts_dur < orig_dur:
                    padded_path = os.path.join(temp_dir, f"padded_timed_{uuid.uuid4().hex}.mp3")
                    try:
                        subprocess.run(
                            [ffmpeg_exe, "-y", "-i", dubbed_audio_path,
                             "-af", f"apad=whole_dur={orig_dur}", "-t", str(orig_dur), padded_path],
                            capture_output=True, check=True
                        )
                        if os.path.exists(padded_path) and os.path.getsize(padded_path) > 0:
                            dubbed_audio_path = padded_path
                    except Exception as pad_e:
                        print(f"Timed pad error: {pad_e}")

                final_audio_path = _mix_bg_with_dub(input_path, dubbed_audio_path, orig_dur, temp_dir, ffmpeg_exe)
                if progress_callback: progress_callback("🎬 Merging mouth-synced dubbed audio with video...")
                return _merge_audio_into_video(input_path, final_audio_path, orig_dur, is_video, temp_dir, progress_callback)

            for _, f in tts_clips:
                if os.path.exists(f): os.remove(f)

    # ── Step 3: Fallback — non-timed full-transcript dubbing ─────────────────
    if progress_callback: progress_callback("⏳ Transcribing speech (fallback mode)...")
    transcript = transcribe_audio_video(input_path)
    if not transcript or "Error" in transcript or len(transcript.strip()) < 3:
        return "ERROR: Could not recognize speech in media."

    if progress_callback: progress_callback(f"🌐 Translating speech to {target_lang_name}...")
    prompt = (
        f"Translate this spoken transcript to natural spoken {target_lang_name}. "
        f"Output ONLY the raw translated text — no headers, no markdown:\n\n{transcript}"
    )
    try:
        try:
            loop = asyncio.get_running_loop()
            translated_text = asyncio.run_coroutine_threadsafe(get_ai_response(0, prompt), loop).result(40)
        except RuntimeError:
            translated_text = asyncio.run(get_ai_response(0, prompt))
    except Exception as e:
        translated_text = transcript

    translated_text = re.sub(r'[*#_~`>\[\]\(\)]', ' ', translated_text)
    translated_text = " ".join(translated_text.split()).strip()
    if not translated_text:
        return "ERROR: Translation resulted in empty speech text."

    if progress_callback: progress_callback("🎙 Generating voiceover dubbing...")
    raw_tts = os.path.join(temp_dir, f"{uuid.uuid4()}_raw_tts.mp3")
    try:
        tts = gTTS(text=translated_text, lang=tts_lang, slow=False)
        tts.save(raw_tts)
        # Studio EQ
        studio_tts = os.path.join(temp_dir, f"{uuid.uuid4()}_studio_tts.mp3")
        try:
            (ffmpeg.input(raw_tts)
             .filter('highpass', f=80).filter('lowpass', f=12000).filter('volume', 1.2)
             .output(studio_tts).overwrite_output()
             .run(capture_output=True, capture_stderr=True))
            if os.path.exists(studio_tts) and os.path.getsize(studio_tts) > 0:
                raw_tts = studio_tts
        except Exception as eq_e:
            print(f"EQ filter error: {eq_e}")
    except Exception as e:
        return f"ERROR: Failed to generate TTS voice: {e}"

    tts_dur = get_video_duration(raw_tts)
    if orig_dur > 0 and tts_dur < orig_dur:
        padded = os.path.join(temp_dir, f"{uuid.uuid4()}_padded.mp3")
        try:
            subprocess.run(
                [ffmpeg_exe, "-y", "-i", raw_tts,
                 "-af", f"apad=whole_dur={orig_dur}", "-t", str(orig_dur), padded],
                capture_output=True, check=True
            )
            if os.path.exists(padded) and os.path.getsize(padded) > 0:
                raw_tts = padded
        except Exception as pad_e:
            print(f"Fallback pad error: {pad_e}")

    final_audio = _mix_bg_with_dub(input_path, raw_tts, orig_dur, temp_dir, ffmpeg_exe)
    return _merge_audio_into_video(input_path, final_audio, orig_dur, is_video, temp_dir, progress_callback)


def _mix_bg_with_dub(input_path, dubbed_path, orig_dur, temp_dir, ffmpeg_exe):
    """Mix original background audio (20%) with dubbed voice (100%) using amix."""
    import subprocess
    orig_audio = os.path.join(temp_dir, f"bg_{uuid.uuid4().hex}.mp3")
    mixed = os.path.join(temp_dir, f"mixed_{uuid.uuid4().hex}.mp3")
    try:
        subprocess.run(
            [ffmpeg_exe, "-y", "-i", input_path, "-vn",
             "-acodec", "libmp3lame", "-q:a", "4", orig_audio],
            capture_output=True, check=True
        )
        if os.path.exists(orig_audio) and os.path.getsize(orig_audio) > 0:
            dur_arg = str(orig_dur) if orig_dur > 0 else "99999"
            subprocess.run(
                [ffmpeg_exe, "-y",
                 "-i", orig_audio, "-i", dubbed_path,
                 "-filter_complex",
                 "[0:a]volume=0.20[bg];[1:a]volume=1.0[voice];[bg][voice]amix=inputs=2:duration=longest:dropout_transition=0[out]",
                 "-map", "[out]", "-t", dur_arg, mixed],
                capture_output=True, check=True
            )
            if os.path.exists(mixed) and os.path.getsize(mixed) > 0:
                if os.path.exists(orig_audio): os.remove(orig_audio)
                return mixed
    except Exception as mix_e:
        print(f"_mix_bg_with_dub error: {mix_e}")
    if os.path.exists(orig_audio): os.remove(orig_audio)
    return dubbed_path  # fallback: no background mix


def _merge_audio_into_video(input_path, audio_path, orig_dur, is_video, temp_dir, progress_callback=None):
    """Merge a final audio track into the video (or return audio for audio-only)."""
    if not is_video:
        return audio_path
    output_path = os.path.join(temp_dir, f"{uuid.uuid4()}_dubbed.mp4")
    try:
        tts_dur = get_video_duration(audio_path)
        if orig_dur > 0 and tts_dur > orig_dur:
            video_in = ffmpeg.input(input_path, stream_loop=-1).video
        else:
            video_in = ffmpeg.input(input_path).video
        audio_in = ffmpeg.input(audio_path).audio
        target_dur = max(orig_dur, tts_dur) if (orig_dur > 0 or tts_dur > 0) else None
        out_opts = {'vcodec': 'copy', 'acodec': 'aac', 'b:a': '192k'}
        if target_dur and target_dur > 0:
            out_opts['t'] = target_dur
        stream = ffmpeg.output(video_in, audio_in, output_path, **out_opts).overwrite_output()
        return _run_ffmpeg_with_progress(stream, output_path, progress_callback)
    except Exception as e:
        print(f"_merge_audio_into_video error: {e}")
        return None

def recap_video_audio(input_path, target_lang='km', is_video=True, voiceover=False, progress_callback=None):
    """
    1. Transcribes full speech in media.
    2. AI generates a SHORT structured Recap/Summary (NOT a full translation).
    3. If voiceover=True, speaks the recap summary once at the start of the video
       with background audio mixed in at 20% volume.
    """
    import asyncio
    import subprocess
    temp_dir = get_temp_dir()

    from plugins.document_parser import transcribe_audio_video
    if progress_callback: progress_callback("⏳ Transcribing media for AI Recap...")
    transcript = transcribe_audio_video(input_path)

    if not transcript or "Error" in transcript or "Unsupported" in transcript or len(transcript.strip()) < 3:
        return "ERROR: Could not extract speech from media to generate recap.", None

    lang_names = {
        'km': 'Khmer', 'en': 'English', 'zh': 'Chinese (Mandarin)',
        'fr': 'French', 'es': 'Spanish', 'ja': 'Japanese',
        'ko': 'Korean', 'de': 'German', 'vi': 'Vietnamese', 'th': 'Thai'
    }
    target_lang_name = lang_names.get(target_lang[:2], 'Khmer')

    if progress_callback: progress_callback(f"🧠 Generating SHORT AI Recap in {target_lang_name}...")

    from plugins.ai_handler import get_ai_response
    # Limit transcript to 8000 chars to avoid AI context overflow
    trimmed_transcript = transcript[:8000]
    prompt = (
        f"You are Udom AI — a professional video summarizer.\n"
        f"Create a SHORT, spoken-style Recap of this video/audio transcript in {target_lang_name}.\n"
        f"CRITICAL RULES:\n"
        f"1. DO NOT translate the full transcript — only summarize it.\n"
        f"2. The recap must be SHORT — 5 to 10 sentences maximum.\n"
        f"3. Write in natural spoken {target_lang_name} that sounds good when read aloud.\n"
        f"4. Include: main topic, key points, and conclusion.\n"
        f"5. Output ONLY the recap text — no headers, no markdown, no bullets.\n\n"
        f"Transcript:\n{trimmed_transcript}"
    )

    try:
        try:
            loop = asyncio.get_running_loop()
            recap_text = asyncio.run_coroutine_threadsafe(get_ai_response(0, prompt), loop).result(45)
        except RuntimeError:
            recap_text = asyncio.run(get_ai_response(0, prompt))
    except Exception as e:
        print(f"AI Recap Error: {e}")
        recap_text = f"Failed to generate recap: {e}"

    if not voiceover:
        return recap_text, None

    # ── Voiceover: speak the recap summary once, play over video ─────────────
    if progress_callback: progress_callback("🎙 Generating Recap voiceover audio...")
    from gtts import gTTS
    tts_lang = 'km' if target_lang.startswith('km') else ('zh-CN' if target_lang.startswith('zh') else target_lang[:2])

    tts_speech = re.sub(r'[*#_~`>\[\]\(\)]', ' ', recap_text)
    tts_speech = re.sub(
        r'^(Here is|Summary|Recap|Khmer|English|Note|Translation):.*$', '',
        tts_speech, flags=re.MULTILINE | re.IGNORECASE
    )
    tts_speech = " ".join(tts_speech.split()).strip()
    if not tts_speech:
        return recap_text, None

    raw_tts = os.path.join(temp_dir, f"{uuid.uuid4()}_recap_raw.mp3")
    try:
        tts = gTTS(text=tts_speech, lang=tts_lang, slow=False)
        tts.save(raw_tts)
        # Studio EQ
        studio_tts = os.path.join(temp_dir, f"{uuid.uuid4()}_recap_studio.mp3")
        try:
            (ffmpeg.input(raw_tts)
             .filter('highpass', f=80).filter('lowpass', f=12000).filter('volume', 1.2)
             .output(studio_tts).overwrite_output()
             .run(capture_stdout=True, capture_stderr=True))
            if os.path.exists(studio_tts) and os.path.getsize(studio_tts) > 0:
                raw_tts = studio_tts
        except Exception as eq_e:
            print(f"Recap EQ error: {eq_e}")
    except Exception as e:
        print(f"Recap TTS Error: {e}")
        return recap_text, None

    ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
    orig_dur = get_video_duration(input_path)
    recap_dur = get_video_duration(raw_tts)

    # Recap voice plays once at start; pad remainder with silence up to orig_dur
    # so the full video plays while background audio continues
    final_recap_audio = raw_tts
    if orig_dur > 0 and recap_dur < orig_dur:
        padded_path = os.path.join(temp_dir, f"{uuid.uuid4()}_recap_padded.mp3")
        try:
            subprocess.run(
                [ffmpeg_exe, "-y", "-i", raw_tts,
                 "-af", f"apad=whole_dur={orig_dur}",
                 "-t", str(orig_dur), padded_path],
                capture_output=True, check=True
            )
            if os.path.exists(padded_path) and os.path.getsize(padded_path) > 0:
                final_recap_audio = padded_path
        except Exception as pad_e:
            print(f"Recap pad error: {pad_e}")

    # Mix: background original audio at 20% + recap voice at 100%
    final_audio = _mix_bg_with_dub(input_path, final_recap_audio, orig_dur, temp_dir, ffmpeg_exe)

    if progress_callback: progress_callback("🎬 Merging recap voiceover with video...")
    res_media = _merge_audio_into_video(input_path, final_audio, orig_dur, is_video, temp_dir, progress_callback)

    return recap_text, res_media

def get_video_duration(file_path):
    """Accurately extracts video or audio duration in seconds using imageio_ffmpeg or ffprobe."""
    try:
        import subprocess
        import re
        ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
        cmd = [ffmpeg_exe, "-i", file_path]
        res = subprocess.run(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE, text=True, encoding="utf-8", errors="ignore")
        match = re.search(r"Duration:\s*(\d+):(\d+):([\d.]+)", res.stderr)
        if match:
            hours, minutes, seconds = match.groups()
            return float(hours) * 3600 + float(minutes) * 60 + float(seconds)
    except Exception as e:
        print(f"Error getting duration with imageio_ffmpeg: {e}")
        
    try:
        probe = ffmpeg.probe(file_path)
        return float(probe['format']['duration'])
    except Exception as e:
        print(f"Error getting video duration: {e}")
        return 0.0

def clip_video_into_parts(input_path, num_clips=3, progress_callback=None):
    """
    Splits a video file into num_clips equal segment files using FFmpeg.
    """
    temp_dir = get_temp_dir()
    duration = get_video_duration(input_path)
    
    if duration <= 0:
        return []
        
    clip_duration = duration / max(1, num_clips)
    output_files = []
    
    for i in range(num_clips):
        start_time = i * clip_duration
        if progress_callback:
            progress_callback(f"✂️ Cutting video clip {i+1} of {num_clips}...")
            
        out_file = os.path.join(temp_dir, f"{uuid.uuid4()}_clip_{i+1}.mp4")
        try:
            stream = (
                ffmpeg
                .input(input_path, ss=start_time, t=clip_duration)
                .output(out_file, vcodec='copy', acodec='copy')
                .overwrite_output()
            )
            ffmpeg.run(stream, capture_stdout=True, capture_stderr=True)
            if os.path.exists(out_file) and os.path.getsize(out_file) > 0:
                output_files.append(out_file)
        except Exception as e:
            print(f"Error creating clip {i+1}: {e}")
            
    return output_files
